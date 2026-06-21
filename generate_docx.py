from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ──────────────────────────────────────────────────────────────────

def border_bottom(paragraph, color='cccccc'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '2')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def spacing(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(13)

def run(p, text, bold=False, italic=False, size=10, rgb=None, name='Calibri'):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = name
    if rgb:
        r.font.color.rgb = RGBColor(*rgb)
    return r

SECONDARY = (110, 110, 110)
TERTIARY  = (150, 150, 150)

def section_title(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=10, after=3)
    r = run(p, text.upper(), bold=True, size=7.5, rgb=SECONDARY)
    r.font.color.rgb = RGBColor(*SECONDARY)
    border_bottom(p)
    return p

def bullet(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=1, after=1)
    pf = p.paragraph_format
    pf.left_indent   = Inches(0.18)
    pf.first_line_indent = Inches(-0.15)
    run(p, '— ', size=8.5, rgb=TERTIARY)
    run(p, text, size=9)
    return p

def job_header(doc, title, company, date, badge=None, desc=None):
    p = doc.add_paragraph()
    spacing(p, before=7, after=0)
    run(p, title, bold=True, size=10)
    run(p, f'  {company}', size=9, rgb=SECONDARY)
    if badge:
        run(p, f'  {badge}', size=8, rgb=(130, 130, 130))
    run(p, f'  {date}', size=8.5, rgb=TERTIARY)
    if desc:
        d = doc.add_paragraph()
        spacing(d, before=1, after=3)
        run(d, desc, size=7.5, italic=True, rgb=(160, 160, 160))
    return p

# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.top_margin      = Inches(1.0)
section.bottom_margin   = Inches(0.75)
section.left_margin     = Inches(0.85)
section.right_margin    = Inches(0.85)
section.header_distance = Inches(0.3)

# Remove default paragraph spacing from Normal style
doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
doc.styles['Normal'].paragraph_format.space_before = Pt(0)

# ── Header ────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
spacing(p, before=10, after=2)
run(p, 'Teerayut Panyoheang', size=26, name='Georgia')

p = doc.add_paragraph()
spacing(p, before=0, after=4)
run(p, 'FULL-STACK DEVELOPER', size=8.5, rgb=SECONDARY)

p = doc.add_paragraph()
spacing(p, before=0, after=0)
run(p, 'Bangkok (relocating)', size=8.5, rgb=SECONDARY)
run(p, '  ·  ', size=8.5, rgb=TERTIARY)
run(p, 'Teerayut.p2710@gmail.com', size=8.5, rgb=SECONDARY)
run(p, '  ·  ', size=8.5, rgb=TERTIARY)
run(p, '(+66) 824 740 430', size=8.5, rgb=SECONDARY)
run(p, '  ·  ', size=8.5, rgb=TERTIARY)
run(p, 'github.com/TeerayutP', size=8.5, rgb=SECONDARY)

# ── Profile ───────────────────────────────────────────────────────────────────

section_title(doc, 'Profile')
p = doc.add_paragraph()
spacing(p, before=0, after=0)
run(p, (
    'Full-stack developer with 5+ years at XENOptics, an R&D company building robotics to automate '
    'high-density fiber optic patching. Started as a frontend engineer and grew into full-stack ownership '
    'as backend responsibilities shifted to me over time — now the sole owner of both sides with no dedicated '
    'backend. Ship across React, Angular, FastAPI, and MySQL in a hardware-software environment, coordinate '
    'across R&D departments, and currently mentor two junior developers. Looking to bring this full-stack depth '
    'into a more structured engineering environment in Bangkok — where I can specialise and go deeper rather than spread thin.'
), size=9, rgb=(80, 80, 80))

# ── Experience ────────────────────────────────────────────────────────────────

section_title(doc, 'Experience')

COMPANY_DESC = 'R&D company developing robotics to automate high-density fiber optic patching  ·  8-person software team  ·  30+ R&D'

job_header(doc, 'Frontend Supervisor', 'XENOptics Ltd.', 'Feb 2025 – Present', badge='↑ Promotion', desc=COMPANY_DESC)
for b in [
    'Own sprint planning and delivery for a 2-person frontend team across 2 live products, maintaining consistent on-time milestone delivery.',
    'Refactored a proprietary monolithic application into a modular architecture.',
    'Delivered a full-stack private portal for role-based fiber management — React frontend, FastAPI APIs, and MySQL data.',
    'Bootstrapped a logistics & shipment web app — set up the React project structure, SPA routing, and reusable component foundation for the team to build on.',
    'Maintained and extended an inherited Docker build pipeline — resolved deployment issues and kept production servers running smoothly.',
]:
    bullet(doc, b)

job_header(doc, 'Frontend Developer', 'XENOptics Ltd.', 'Sep 2020 – Feb 2025', desc=COMPANY_DESC)
for b in [
    'Built and maintained an Angular SPA for fiber patching software (Angular Material), serving a network dashboard managing 2,000+ fiber connections and devices.',
    'Developed a React-based NMS (Network Management System) SPA with HeroUI and Tailwind customization.',
    'Customised 15+ HeroUI components with a unified theme — structured as an atomic component library reused across 2 products, ensuring consistent UI with minimal duplication.',
    'Took on full backend ownership during a period without a dedicated backend developer — maintained a legacy PHP codebase and a Lumen API to keep production running without interruption.',
    'Designed and implemented RESTful APIs with FastAPI: JWT auth, role-based permissions, and business logic in MySQL stored procedures.',
    'Collaborated closely with the hardware production team — adapted frontend features to meet evolving hardware requirements and kept the software pipeline aligned with physical production workflows.',
    'Contributed Python code for robot machine-state management; built an internal web interface for debugging and configuring robotic arms.',
    'Mentored and trained junior frontend developers on team practices, code standards, and best practices.',
]:
    bullet(doc, b)

job_header(doc, 'Web Development Intern', 'Chiang Mai Air Traffic Control Centre', 'Jun 2019 – Jul 2019', badge='Internship')
bullet(doc, 'Engineer Relocation Planner — Designed and built a web application for engineers to plan and visualise a provincial centre relocation.')

# ── Education ─────────────────────────────────────────────────────────────────

section_title(doc, 'Education')

p = doc.add_paragraph()
spacing(p, before=4, after=1)
run(p, 'Bachelor of Engineering in Computer Engineering', bold=True, size=10)
run(p, '  Aug 2016 – May 2020', size=8.5, rgb=TERTIARY)

p = doc.add_paragraph()
spacing(p, before=0, after=4)
run(p, 'Chiang Mai University  ·  Chiang Mai, Thailand', size=9, rgb=SECONDARY)

p = doc.add_paragraph()
spacing(p, before=2, after=1)
run(p, 'Indoor Surveillance System', bold=True, size=9.5)
run(p, '  Senior Thesis  ·  2020', size=8, rgb=SECONDARY)

for b in [
    'Built a web application for indoor positioning using Wi-Fi RSSI signal log processing with Node.js and Firebase.',
    'Rendered interactive indoor maps with LeafletJS to visualise real-time device locations.',
]:
    bullet(doc, b)

# ── Skills ────────────────────────────────────────────────────────────────────

section_title(doc, 'Skills')

for label, tags in [
    ('Frontend', 'React · Angular · TypeScript · Redux · HeroUI · Tailwind · SCSS'),
    ('Backend',  'FastAPI · Python · PHP/Lumen · MySQL · SQLite'),
    ('DevOps',   'Docker · Git · Vercel · Railway · Vite · Webpack · npm'),
]:
    p = doc.add_paragraph()
    spacing(p, before=2, after=2)
    run(p, label.upper() + '   ', bold=True, size=7.5, rgb=SECONDARY)
    run(p, tags, size=9)

# ── Projects ──────────────────────────────────────────────────────────────────

section_title(doc, 'Projects')

p = doc.add_paragraph()
spacing(p, before=4, after=1)
run(p, 'Bookmarker', bold=True, size=10)
run(p, '  bookmarker-dun-three.vercel.app', size=9, rgb=SECONDARY)
bullet(doc, 'Personal project — deployed on Vercel with backend on Railway.')

# ── Additional ────────────────────────────────────────────────────────────────

section_title(doc, 'Additional')

p = doc.add_paragraph()
spacing(p, before=4, after=0)
run(p, 'NATIONALITY   ', bold=True, size=7.5, rgb=SECONDARY)
run(p, 'Thai', size=9)

p = doc.add_paragraph()
spacing(p, before=2, after=0)
run(p, 'LANGUAGES      ', bold=True, size=7.5, rgb=SECONDARY)
run(p, 'Thai (native)  ·  English (working — written communication and technical discussions)', size=9)

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save('Teerayut_CV.docx')
print('Done: Teerayut_CV.docx')
